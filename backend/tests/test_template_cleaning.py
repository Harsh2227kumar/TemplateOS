"""
Data-layer tests for manual template cleaning.
V1.3 Phase 2 — Member 3

Layers:
1. Schema validation (run now): PlaceholderReplacement key/type rules, the
   confirm gate, min-length on replacements.
2. CRUD unit tests (run now): set_processed_path, advance_status
   (forward-only), next_display_order, append_field (idempotent, no dup).
3. Cleaner unit tests (run now): apply_replacements on a real in-memory DOCX
   built with python-docx — asserts {{key}} produced, original bytes
   unchanged, and paragraph-level rewrite across split runs.
4. Endpoint tests (skip until Member 2's POST /clean + GET /content routes
   exist): owner/confirm guards, processed path + field creation with
   example_value = sample text, status advance, re-clean idempotency.
"""

import io
import os
import zipfile

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

os.environ.setdefault("DATABASE_URL", "sqlite://")
os.environ.setdefault(
    "JWT_SECRET_KEY", "test-secret-that-is-long-enough-for-auth-tests"
)

from app.crud.template_crud import advance_status, set_processed_path
from app.crud.template_field_crud import (
    append_field,
    bulk_create_fields,
    next_display_order,
)
from app.db.base import Base
from app.db.session import get_db
from app.main import app
from app.models.template import Template
from app.models.user import User  # noqa: F401 — registers the users table
from app.schemas.cleaning import (
    CleanTemplateRequest,
    PlaceholderReplacement,
)
from app.schemas.template_field import TemplateFieldCreate

test_engine = create_engine(
    "sqlite://",
    connect_args={"check_same_thread": False},
    poolclass=StaticPool,
)
TestSession = sessionmaker(bind=test_engine, autoflush=False, autocommit=False)


def override_get_db():
    with TestSession() as db:
        yield db


client = TestClient(app)


@pytest.fixture(scope="module", autouse=True)
def setup_db():
    old = app.dependency_overrides.get(get_db)
    app.dependency_overrides[get_db] = override_get_db
    Base.metadata.drop_all(bind=test_engine)
    Base.metadata.create_all(bind=test_engine)
    yield
    Base.metadata.drop_all(bind=test_engine)
    if old is not None:
        app.dependency_overrides[get_db] = old
    else:
        app.dependency_overrides.pop(get_db, None)


@pytest.fixture()
def db():
    """Function-scoped session for direct CRUD tests (isolated per test)."""
    session = TestSession()
    yield session
    session.rollback()
    session.close()


def _make_template(
    session, name: str = "Cleaning Test", status: str = "uploaded"
) -> Template:
    template = Template(
        name=name,
        category="mom",
        visibility="private",
        uploaded_by=1,
        status=status,
    )
    session.add(template)
    session.commit()
    session.refresh(template)
    return template


# ── 1. Schema validation ──────────────────────────────────────────


def test_replacement_accepts_valid_key_and_type():
    replacement = PlaceholderReplacement(
        sample_text="Annual Fest Planning Meeting",
        placeholder_key="meeting_title",
        field_type="textarea",
    )
    assert replacement.placeholder_key == "meeting_title"
    assert replacement.field_type == "textarea"


def test_replacement_rejects_invalid_keys():
    for bad_key in (
        "Meeting Title",
        "2date",
        "camelCase",
        "has-dash",
        "has.dot",
        "",
        "UPPER",
    ):
        with pytest.raises(Exception):
            PlaceholderReplacement(sample_text="x", placeholder_key=bad_key)


def test_replacement_rejects_invalid_field_type():
    with pytest.raises(Exception):
        PlaceholderReplacement(
            sample_text="x", placeholder_key="ok_key", field_type="dropdown"
        )


def test_replacement_requires_sample_text():
    with pytest.raises(Exception):
        PlaceholderReplacement(sample_text="", placeholder_key="ok_key")


def test_clean_request_requires_replacements():
    with pytest.raises(Exception):
        CleanTemplateRequest(replacements=[])


def test_clean_request_confirm_defaults_false():
    """The confirm gate defaults to false — the endpoint must reject unconfirmed cleans."""
    request = CleanTemplateRequest(
        replacements=[PlaceholderReplacement(sample_text="x", placeholder_key="ok_key")]
    )
    assert request.confirm is False
    assert request.mark_configured is False


# ── 2. CRUD helpers ───────────────────────────────────────────────


def test_set_processed_path_records_path_and_preserves_original(db):
    template = _make_template(db)
    template.original_file_path = "templates/original/keep-me.docx"
    db.commit()

    updated = set_processed_path(db, template, "templates/processed/new-me.docx")

    assert updated.processed_file_path == "templates/processed/new-me.docx"
    assert updated.original_file_path == "templates/original/keep-me.docx"


def test_advance_status_moves_forward(db):
    template = _make_template(db, status="uploaded")
    updated = advance_status(db, template, "placeholder_detected")
    assert updated.status == "placeholder_detected"
    updated = advance_status(db, template, "field_configured")
    assert updated.status == "field_configured"


def test_advance_status_never_downgrades(db):
    template = _make_template(db, status="field_configured")
    updated = advance_status(db, template, "placeholder_detected")
    assert updated.status == "field_configured"  # unchanged

    template = _make_template(db, name="Active Template", status="active")
    updated = advance_status(db, template, "field_configured")
    assert updated.status == "active"  # unchanged


def test_advance_status_rejects_unknown_status(db):
    template = _make_template(db)
    with pytest.raises(ValueError):
        advance_status(db, template, "not_a_status")


def test_next_display_order_empty_then_increments(db):
    template = _make_template(db)
    assert next_display_order(db, template.id) == 0
    append_field(db, template.id, TemplateFieldCreate(field_name="first"))
    assert next_display_order(db, template.id) == 1
    append_field(db, template.id, TemplateFieldCreate(field_name="second"))
    assert next_display_order(db, template.id) == 2


def test_append_field_skips_duplicates(db):
    template = _make_template(db)
    first = append_field(
        db, template.id, TemplateFieldCreate(field_name="meeting_title")
    )
    assert first is not None
    duplicate = append_field(
        db, template.id, TemplateFieldCreate(field_name="meeting_title")
    )
    assert duplicate is None  # skipped, not created
    from app.crud.template_field_crud import get_fields_by_template

    assert len(get_fields_by_template(db, template.id)) == 1


def test_append_field_assigns_next_order_automatically(db):
    template = _make_template(db)
    bulk_create_fields(
        db,
        template.id,
        [
            TemplateFieldCreate(field_name="a", display_order=0),
            TemplateFieldCreate(field_name="b", display_order=1),
        ],
    )
    appended = append_field(
        db,
        template.id,
        TemplateFieldCreate(field_name="meeting_title", source="cleaned"),
    )
    assert appended is not None
    assert appended.display_order == 2
    assert appended.source == "cleaned"


# ── 3. Cleaner unit tests (Member 2's service — skip until it ships) ──

try:
    from app.services.docx_cleaner import apply_replacements, ReplacementSpec  # noqa: F401

    DOCX_CLEANER_AVAILABLE = True
except ModuleNotFoundError:
    DOCX_CLEANER_AVAILABLE = False

requires_docx_cleaner = pytest.mark.skipif(
    not DOCX_CLEANER_AVAILABLE,
    reason="app/services/docx_cleaner.py not implemented yet (Member 2, V1.3 Phase 2)",
)


def _build_sample_docx() -> bytes:
    """Real DOCX via python-docx: normal paragraph + one split across two runs."""
    from docx import Document

    document = Document()
    document.add_paragraph("Subject: Annual Fest Planning Meeting")
    # Split the sample text across two runs inside one paragraph.
    paragraph = document.add_paragraph()
    paragraph.add_run("Meeting: Annual Fest ")
    paragraph.add_run("Planning Meeting")
    document.add_paragraph("Regards, Secretary")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@requires_docx_cleaner
def test_apply_replacements_produces_placeholder_and_preserves_original():
    from app.services.docx_cleaner import apply_replacements

    def _document_xml(docx: bytes) -> bytes:
        """A DOCX is a zip; check word/document.xml (uncompressed via zipfile)."""
        with zipfile.ZipFile(io.BytesIO(docx)) as archive:
            return archive.read("word/document.xml")

    original_bytes = _build_sample_docx()
    processed = apply_replacements(
        original_bytes,
        [
            {
                "sample_text": "Annual Fest Planning Meeting",
                "placeholder_key": "meeting_title",
            }
        ],
    )
    assert processed != original_bytes
    assert b"{{ meeting_title }}" in _document_xml(
        processed
    ) or b"{{meeting_title}}" in _document_xml(processed)
    # Original is untouched: bytes identical, no placeholder introduced.
    assert original_bytes == _build_sample_docx()
    assert b"{{ meeting_title }}" not in _document_xml(original_bytes)


@requires_docx_cleaner
def test_apply_replacements_handles_split_runs_and_counts_occurrences():
    from app.services.docx_cleaner import apply_replacements, ReplacementSpec

    original_bytes = _build_sample_docx()
    processed = apply_replacements(
        original_bytes,
        [
            ReplacementSpec(
                sample_text="Annual Fest Planning Meeting",
                placeholder_key="meeting_title",
            )
        ],
    )
    from docx import Document

    doc = Document(io.BytesIO(processed))
    texts = [p.text for p in doc.paragraphs]
    # Both occurrences replaced — including the one split across two runs.
    assert texts[0] == "Subject: {{ meeting_title }}"
    assert texts[1] == "Meeting: {{ meeting_title }}"


@requires_docx_cleaner
def test_apply_replacements_reports_unmatched():
    from app.services.docx_cleaner import apply_replacements, ReplacementSpec

    original_bytes = _build_sample_docx()
    processed, unmatched = apply_replacements(
        original_bytes,
        [
            ReplacementSpec(sample_text="No Such Text", placeholder_key="ghost_key"),
            ReplacementSpec(sample_text="Secretary", placeholder_key="sender"),
        ],
        return_unmatched=True,
    )
    assert unmatched == ["No Such Text"]
    from docx import Document

    doc = Document(io.BytesIO(processed))
    assert "{{ sender }}" in doc.paragraphs[2].text


# ── 4. Endpoint tests (skip until Member 2's clean routes exist) ────

ROUTE_PATHS = {getattr(route, "path", "") for route in app.routes}
CLEAN_ENDPOINT_MISSING = not any(p.endswith("/clean") for p in ROUTE_PATHS)

requires_clean_endpoint = pytest.mark.skipif(
    CLEAN_ENDPOINT_MISSING,
    reason="POST /templates/{id}/clean not implemented yet (Member 2, V1.3 Phase 2)",
)


def _signup_and_login(email: str, full_name: str) -> str:
    client.post(
        "/api/v1/auth/signup",
        json={"full_name": full_name, "email": email, "password": "strong-password"},
    )
    login = client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": "strong-password"},
    )
    return login.json()["access_token"]


@pytest.fixture(scope="module")
def clean_owner_token():
    return _signup_and_login("clean_owner@example.com", "Clean Owner")


@pytest.fixture(scope="module")
def clean_other_token():
    return _signup_and_login("clean_other@example.com", "Clean Other")


def _upload_template(token: str, name: str) -> dict:
    """Upload a real sample DOCX through the API (storage write is real, local disk)."""
    files = {
        "file": (
            "sample.docx",
            io.BytesIO(_build_sample_docx()),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    data = {"name": name, "category": "mom", "visibility": "private"}
    response = client.post(
        "/api/v1/templates/upload",
        files=files,
        data=data,
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 201, f"Upload failed: {response.json()}"
    return response.json()


@requires_clean_endpoint
def test_clean_persists_processed_path_and_field(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean Happy Path")
    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
        json={
            "replacements": [
                {
                    "sample_text": "Annual Fest Planning Meeting",
                    "placeholder_key": "meeting_title",
                }
            ],
            "confirm": True,
        },
    )
    assert response.status_code == 200, response.json()
    body = response.json()

    assert body["template_id"] == template["id"]
    assert body["processed_file_path"]
    assert body["status"] == "placeholder_detected"

    created = body["created_fields"]
    assert len(created) == 1
    assert created[0]["field_name"] == "meeting_title"
    assert created[0]["example_value"] == "Annual Fest Planning Meeting"
    assert created[0]["source"] == "cleaned"

    results = body["results"]
    assert len(results) == 1
    assert results[0]["matched"] is True
    assert results[0]["occurrences"] == 2  # body + split-run paragraph
    assert body["warnings"] == {"unmatched": [], "invalid_keys": []}

    # Original preserved on the record; processed path set.
    detail = client.get(
        f"/api/v1/templates/{template['id']}",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
    ).json()
    assert detail["original_file_path"] == template["original_file_path"]
    assert detail["processed_file_path"] == body["processed_file_path"]

    # The processed DOCX on disk really contains the placeholder
    # (read word/document.xml from the zip — entries are DEFLATE-compressed).
    from app.services.storage_service import storage_service

    processed_bytes = storage_service.read_bytes(body["processed_file_path"])
    with zipfile.ZipFile(io.BytesIO(processed_bytes)) as archive:
        document_xml = archive.read("word/document.xml")
    assert (
        b"{{ meeting_title }}" in document_xml or b"{{meeting_title}}" in document_xml
    )


@requires_clean_endpoint
def test_clean_requires_confirm(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean No Confirm")
    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
        json={
            "replacements": [{"sample_text": "Secretary", "placeholder_key": "sender"}],
            "confirm": False,
        },
    )
    assert response.status_code == 400
    # Nothing was persisted.
    detail = client.get(
        f"/api/v1/templates/{template['id']}",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
    ).json()
    assert detail["processed_file_path"] is None


@requires_clean_endpoint
def test_clean_rejects_non_owner(clean_owner_token, clean_other_token):
    template = _upload_template(clean_owner_token, "Clean Ownership")
    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_other_token}"},
        json={
            "replacements": [{"sample_text": "Secretary", "placeholder_key": "sender"}],
            "confirm": True,
        },
    )
    assert response.status_code == 403


@requires_clean_endpoint
def test_clean_rejects_empty_replacements(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean Empty")
    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
        json={"replacements": [], "confirm": True},
    )
    assert response.status_code in (400, 422)


@requires_clean_endpoint
def test_clean_rejects_invalid_key_via_schema(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean Invalid Key")
    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
        json={
            "replacements": [
                {"sample_text": "Secretary", "placeholder_key": "Bad Key"}
            ],
            "confirm": True,
        },
    )
    assert response.status_code == 422  # schema rejects before reaching the endpoint


@requires_clean_endpoint
def test_clean_is_idempotent_across_reruns(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean Idempotent")
    headers = {"Authorization": f"Bearer {clean_owner_token}"}
    payload = {
        "replacements": [
            {
                "sample_text": "Annual Fest Planning Meeting",
                "placeholder_key": "meeting_title",
            }
        ],
        "confirm": True,
    }
    first = client.post(
        f"/api/v1/templates/{template['id']}/clean", headers=headers, json=payload
    )
    assert first.status_code == 200
    second = client.post(
        f"/api/v1/templates/{template['id']}/clean", headers=headers, json=payload
    )
    assert second.status_code == 200

    # Field exists once, not duplicated; processed path refreshed.
    fields = client.get(
        f"/api/v1/templates/{template['id']}/fields",
        headers=headers,
    ).json()
    meeting_title_rows = [f for f in fields if f["field_name"] == "meeting_title"]
    assert len(meeting_title_rows) == 1
    assert second.json()["processed_file_path"]


@requires_clean_endpoint
def test_clean_mark_configured_advances_status(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean Mark Configured")
    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
        json={
            "replacements": [{"sample_text": "Secretary", "placeholder_key": "sender"}],
            "confirm": True,
            "mark_configured": True,
        },
    )
    assert response.status_code == 200
    assert response.json()["status"] == "field_configured"


@requires_clean_endpoint
def test_clean_does_not_downgrade_configured_status(clean_owner_token):
    template = _upload_template(clean_owner_token, "Clean No Downgrade")
    with TestSession() as session:
        record = session.get(Template, template["id"])
        assert record is not None, "template must exist"
        record.status = "field_configured"
        session.commit()

    response = client.post(
        f"/api/v1/templates/{template['id']}/clean",
        headers={"Authorization": f"Bearer {clean_owner_token}"},
        json={
            "replacements": [{"sample_text": "Secretary", "placeholder_key": "sender"}],
            "confirm": True,
        },
    )
    assert response.status_code == 200
    assert response.json()["status"] == "field_configured"  # stayed, not downgraded
