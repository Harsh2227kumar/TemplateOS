"""
DOCX cleaning service (V1.3 Phase 2 — Member 2).

Pure replacement engine — NO database access. Turns owner-confirmed sample
text into {{ placeholders }} across body paragraphs, table cells (recursing
into nested tables), and section headers/footers, producing a NEW processed
DOCX. The caller's bytes are never mutated: the engine works on a copy.

Hard rules (MD/features.md §4, MD/mvp-scope.md):
- Manual confirmation only — exactly the replacements the owner confirmed.
- The original file is preserved; this module never writes to storage.
- Placeholder keys must pass Phase 1's VALID_KEY_PATTERN; invalid keys are
  skipped and reported, never injected into the document.
"""

from dataclasses import dataclass
from io import BytesIO
from typing import Literal, overload

from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from app.services.docx_parser import VALID_KEY_PATTERN

# Accepted MVP tradeoff (features.md: "formatting-aware cleaning" is a FUTURE
# item): a rewritten paragraph is normalized to its first run's formatting.


@dataclass(frozen=True)
class ReplacementSpec:
    """One confirmed sample_text -> placeholder_key conversion."""

    sample_text: str
    placeholder_key: str
    location_hint: str | None = None
    segment_index: int | None = None


@dataclass(frozen=True)
class ReplacementResult:
    """Per-replacement outcome (mirrors the schema Member 3 defined)."""

    placeholder_key: str
    sample_text: str
    occurrences: int = 0
    matched: bool = False
    reason: str | None = None


def _replace_in_paragraph(paragraph: Paragraph, find: str, repl: str) -> int:
    """
    Replace every exact occurrence of `find` in the paragraph with `repl`.

    Word splits one visible string across multiple runs, so matching happens
    on the paragraph's FULL text and the runs are rewritten: the first run
    keeps its formatting and receives the whole new text; the remaining runs
    are blanked. Returns the number of occurrences replaced.
    """
    full = paragraph.text
    if not find or find not in full:
        return 0
    new_full = full.replace(find, repl)
    if paragraph.runs:
        paragraph.runs[0].text = new_full
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_full
    return full.count(find)


def _iter_paragraphs(document: DocumentObject):
    """Yield every paragraph in body (document order), table cells (recursive), headers, footers."""
    for paragraph in document.paragraphs:
        yield paragraph

    def walk_table(table: Table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    yield from walk_table(nested)

    for table in document.tables:
        yield from walk_table(table)

    for section in document.sections:
        header_footer_parts = (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        )
        for part in header_footer_parts:
            if part is None or part.is_linked_to_previous:
                continue
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                yield from walk_table(table)


def _normalize_specs(replacements) -> list[tuple[ReplacementSpec, bool]]:
    """Accept ReplacementSpec objects, Pydantic models, or raw dicts; validate keys against VALID_KEY_PATTERN."""
    specs: list[tuple[ReplacementSpec, bool]] = []
    for item in replacements:
        if isinstance(item, ReplacementSpec):
            spec = item
        elif isinstance(item, dict):
            spec = ReplacementSpec(
                sample_text=item["sample_text"],
                placeholder_key=item["placeholder_key"],
                location_hint=item.get("location_hint"),
                segment_index=item.get("segment_index"),
            )
        else:  # Pydantic model (e.g. schemas.cleaning.PlaceholderReplacement)
            spec = ReplacementSpec(
                sample_text=item.sample_text,
                placeholder_key=item.placeholder_key,
                location_hint=getattr(item, "location_hint", None),
                segment_index=getattr(item, "segment_index", None),
            )
        is_valid = bool(VALID_KEY_PATTERN.match(spec.placeholder_key))
        specs.append((spec, is_valid))
    return specs


@overload
def apply_replacements(
    docx_bytes: bytes, replacements, return_unmatched: Literal[False] = False
) -> bytes: ...


@overload
def apply_replacements(
    docx_bytes: bytes, replacements, return_unmatched: Literal[True]
) -> tuple[bytes, list[str]]: ...


def apply_replacements(docx_bytes: bytes, replacements, return_unmatched: bool = False):
    """
    Apply confirmed replacements to a COPY of the DOCX and return the processed bytes.

    Returns:
        processed_bytes
            when return_unmatched is False (default)
        (processed_bytes, unmatched_sample_texts)
            when return_unmatched is True — sample texts found 0 times

    The endpoint uses apply_replacements_with_results() for the full per-
    replacement outcomes.
    """
    processed, results = _apply(docx_bytes, replacements)
    if return_unmatched:
        unmatched = [r.sample_text for r in results if not r.matched]
        return processed, unmatched
    return processed


def apply_replacements_with_results(docx_bytes: bytes, replacements):
    """
    Same as apply_replacements, but returns the full per-replacement results:
    (processed_bytes, list[ReplacementResult]) — occurrences, matched, reason.
    """
    return _apply(docx_bytes, replacements)


def _apply(docx_bytes: bytes, replacements) -> tuple[bytes, list[ReplacementResult]]:
    """Core engine: validate keys, rewrite paragraphs, return (bytes, results)."""
    # Never mutate the caller's buffer: work on an in-memory copy.
    tpl = DocxTemplate(BytesIO(docx_bytes))
    tpl.init_docx()
    document = tpl.docx
    assert document is not None, "docxtpl failed to initialize the document"

    paragraphs = list(_iter_paragraphs(document))

    results: list[ReplacementResult] = []
    for spec, is_valid in _normalize_specs(replacements):
        if not is_valid:
            results.append(
                ReplacementResult(
                    placeholder_key=spec.placeholder_key,
                    sample_text=spec.sample_text,
                    occurrences=0,
                    matched=False,
                    reason="invalid_key",
                )
            )
            continue
        replacement_text = "{{ " + spec.placeholder_key + " }}"
        occurrences = 0
        for paragraph in paragraphs:
            occurrences += _replace_in_paragraph(
                paragraph, spec.sample_text, replacement_text
            )
        results.append(
            ReplacementResult(
                placeholder_key=spec.placeholder_key,
                sample_text=spec.sample_text,
                occurrences=occurrences,
                matched=occurrences > 0,
                reason=None if occurrences > 0 else "sample_text_not_found",
            )
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), results


def get_renderable_path(template) -> str:
    """
    Generation-source contract (V1.6): the fill-ready file is the processed
    copy when it exists, else the original. Returns '' when neither is set.
    """
    return template.processed_file_path or template.original_file_path or ""
