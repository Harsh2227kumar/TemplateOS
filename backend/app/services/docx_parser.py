"""
DOCX placeholder detection service (V1.3 Phase 1 — Member 2).

Pure parsing — NO database access. The detection endpoint
(app/api/v1/endpoints/templates.py) persists through Member 3's
template_field_crud using this module's DetectionResult.

Design (detect/generate parity):
- docxtpl (the SAME Jinja2 engine that renders documents in V1.6) provides
  the AUTHORITATIVE placeholder set via get_undeclared_template_variables().
- A permissive regex scan over extracted text provides ordering, duplicate
  counts, and malformed-token warnings ({{Bad Name}}) that Jinja cannot parse.
"""

import itertools
import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import TypedDict

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

# Permissive on purpose: must ALSO catch invalid names like "{{Meeting Title}}"
# so the UI can warn (docxtpl/Jinja2 reject those outright).
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

# A valid placeholder key: lowercase snake_case, starts with a letter.
VALID_KEY_PATTERN = re.compile(r"^[a-z][a-z0-9_]*$")

_INVALID_NAME_REASON = (
    "Placeholder names must be lowercase snake_case "
    "(letters, digits, underscores) and start with a letter"
)


class TextSegment(TypedDict):
    index: int
    location: str  # "body" | "table" | "header" | "footer"
    text: str


@dataclass(frozen=True)
class DetectedField:
    """A valid placeholder key, in first-seen document order."""

    key: str
    display_order: int


@dataclass(frozen=True)
class InvalidNameWarning:
    raw: str
    suggested_key: str
    count: int
    reason: str


@dataclass(frozen=True)
class DuplicateWarning:
    key: str
    count: int


@dataclass(frozen=True)
class DetectionResult:
    valid_fields: list[DetectedField] = field(default_factory=list)
    invalid_names: list[InvalidNameWarning] = field(default_factory=list)
    duplicates: list[DuplicateWarning] = field(default_factory=list)
    total_matches: int = 0
    unique_valid: int = 0
    parse_warning: str | None = None


def suggest_key(raw: str) -> str:
    """Normalize a malformed token into a suggested snake_case key."""
    key = raw.strip().lower()
    key = re.sub(r"[\s\-]+", "_", key)
    key = re.sub(r"[^a-z0-9_]", "", key)
    key = re.sub(r"_{2,}", "_", key).strip("_")
    if not key:
        return "field"
    if not key[0].isalpha():
        key = f"field_{key}"
    return key


def humanize_key(key: str) -> str:
    """'employee_name' -> 'Employee Name' (field_label for detected fields)."""
    return " ".join(part.capitalize() for part in key.split("_") if part)


def extract_text_segments(docx_bytes: bytes) -> list[TextSegment]:
    """
    Walk ALL text locations of the DOCX in document order and return ordered
    segments: body paragraphs, table cells (recursing into nested tables),
    and the headers/footers of each section.

    Reuses one DocxTemplate parse (tpl.docx) so detection and extraction share
    a single document open. Reused by Phase 2 (cleaning) and Phase 4 (AI
    context excerpt). python-docx's paragraph.text already joins a paragraph's
    runs, so a placeholder split across runs INSIDE one paragraph reads back
    whole here.
    """
    tpl = DocxTemplate(BytesIO(docx_bytes))
    tpl.init_docx()
    document = tpl.docx
    assert document is not None, "docxtpl failed to initialize the document"

    segments: list[TextSegment] = []
    counter = itertools.count()

    def add_text(text: str, location: str) -> None:
        if text and text.strip():
            segments.append(
                TextSegment(index=next(counter), location=location, text=text)
            )

    def walk_table(table: Table, location: str) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    add_text(paragraph.text, location)
                for nested_table in cell.tables:
                    walk_table(nested_table, location)

    # Body in true document order (top-level paragraphs and tables interleaved).
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            add_text(Paragraph(child, document).text, "body")
        elif child.tag == qn("w:tbl"):
            walk_table(Table(child, document), "table")

    # Headers and footers of each section (skip inherited/linked definitions).
    for section in document.sections:
        header_footer_parts = (
            (section.header, "header"),
            (section.footer, "footer"),
            (section.first_page_header, "header"),
            (section.first_page_footer, "footer"),
            (section.even_page_header, "header"),
            (section.even_page_footer, "footer"),
        )
        for part, location in header_footer_parts:
            if part is None or part.is_linked_to_previous:
                continue
            for paragraph in part.paragraphs:
                add_text(paragraph.text, location)
            for table in part.tables:
                walk_table(table, location)

    return segments


def get_template_variables(docx_bytes: bytes) -> tuple[set[str] | None, str | None]:
    """
    THE AUTHORITATIVE detector: docxtpl/Jinja2, the same engine that renders
    documents in V1.6 — so what we detect is exactly what will fill later.

    Returns (variables, None) on success. Jinja2 cannot parse malformed names
    (e.g. {{Meeting Title}}, {{2date}}): the call raises and we return
    (None, message) — never let it 500; the regex layer still surfaces those
    tokens as invalid-name warnings.
    """
    try:
        tpl = DocxTemplate(BytesIO(docx_bytes))
        variables = tpl.get_undeclared_template_variables()
        return set(variables), None
    except Exception as exc:  # jinja2.TemplateSyntaxError, docxtpl errors, ...
        return None, f"The template could not be parsed for document generation: {exc}"


def detect_placeholders(docx_bytes: bytes) -> DetectionResult:
    """
    Detect every {{placeholder}} in the DOCX. Read-only — never modifies the file.

    Step 1 (regex): scan extracted text segments for ordering, duplicate counts,
    and malformed tokens.
    Step 2 (docxtpl): intersect with the authoritative variable set for
    detect/generate parity; on parse failure fall back to the regex key-valid
    tokens and surface a parse warning.
    """
    # Step 1 — warnings + ordering.
    segments = extract_text_segments(docx_bytes)
    raw_info: dict[str, dict] = {}
    total_matches = 0
    for segment in segments:
        for match in PLACEHOLDER_PATTERN.finditer(segment["text"]):
            raw = match.group(1)
            total_matches += 1
            info = raw_info.setdefault(
                raw, {"count": 0, "first_order": total_matches - 1}
            )
            info["count"] += 1

    valid_keys: list[str] = []
    invalid_names: list[InvalidNameWarning] = []
    for raw, info in raw_info.items():
        if VALID_KEY_PATTERN.match(raw):
            valid_keys.append(raw)
        else:
            invalid_names.append(
                InvalidNameWarning(
                    raw=raw,
                    suggested_key=suggest_key(raw),
                    count=info["count"],
                    reason=_INVALID_NAME_REASON,
                )
            )

    # Step 2 — authority + parity.
    variables, parse_warning = get_template_variables(docx_bytes)
    if variables is not None:
        # Persist only tokens that are key-valid AND seen by docxtpl — they
        # are guaranteed to render in V1.6. docxtpl-only variables the regex
        # missed (rare, e.g. inside a text box) are appended after.
        persisted_keys = [raw for raw in valid_keys if raw in variables]
        regex_seen = set(valid_keys)
        extras = sorted(
            variable
            for variable in variables
            if variable not in regex_seen and VALID_KEY_PATTERN.match(variable)
        )
        persisted_keys.extend(extras)
    else:
        # docxtpl could not parse the template (malformed tokens): fall back
        # to the regex key-valid tokens so detection still succeeds. The UI
        # shows parse_warning so the owner can fix the tokens before generation.
        persisted_keys = list(valid_keys)

    # Order by first appearance; docxtpl-only extras come after regex-seen keys.
    first_order = {raw: info["first_order"] for raw, info in raw_info.items()}
    persisted_keys.sort(key=lambda key: first_order.get(key, len(first_order) + 1))
    valid_fields = [
        DetectedField(key=key, display_order=order)
        for order, key in enumerate(persisted_keys)
    ]

    duplicates = [
        DuplicateWarning(key=raw, count=info["count"])
        for raw, info in raw_info.items()
        if info["count"] > 1
    ]

    return DetectionResult(
        valid_fields=valid_fields,
        invalid_names=invalid_names,
        duplicates=duplicates,
        total_matches=total_matches,
        unique_valid=len(valid_fields),
        parse_warning=parse_warning,
    )
